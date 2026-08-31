import os
import pypdf
import io
from PIL import Image
from typing import Callable, Dict
import fitz  # PyMuPDF

class EncryptedPDFError(Exception):
    """自訂例外：當 PDF 受密碼保護且未能解密時拋出"""
    def __init__(self, filepath: str):
        self.filepath = filepath
        super().__init__(f"檔案受密碼保護: {filepath}")

def merge_pdfs(pdf_list: list[str], output_folder: str, output_filename: str, callback: Callable[[float], None] | None = None) -> tuple[bool, str]:
    """
    合併多個 PDF 檔案。
    """
    try:
        writer = pypdf.PdfWriter()
        total = len(pdf_list)
        
        if not output_filename.lower().endswith(".pdf"):
            output_filename += ".pdf"
            
        output_path = os.path.join(output_folder, output_filename)

        if not os.path.exists(output_folder):
            os.makedirs(output_folder, exist_ok=True)

        for i, pdf in enumerate(pdf_list):
            writer.append(pdf)
            if callback:
                callback((i + 1) / total)
        
        with open(output_path, "wb") as f:
            writer.write(f)
            
        return True, f"合併成功！檔案已儲存至：\n{output_path}"
    except Exception as e:
        return False, f"合併失敗: {str(e)}"
    finally:
        try:
            writer.close()
        except:
            pass

def parse_range_string(ranges: str, total_pages: int) -> list[int]:
    """
    解析範圍字串（如 "1-5, 8, 10-12"）並回傳 0-indexed 的頁碼列表。
    """
    target_pages = set()
    if not ranges.strip():
        return list(range(total_pages))
        
    parts = ranges.replace(" ", "").split(",")
    for part in parts:
        try:
            if "-" in part:
                start_str, end_str = part.split("-")
                s = int(start_str)
                e = int(end_str)
                for p in range(s, e + 1):
                    if 1 <= p <= total_pages:
                        target_pages.add(p - 1)
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    target_pages.add(p - 1)
        except ValueError:
            continue # 忽略格式錯誤的部分
            
    return sorted(list(target_pages))

def split_pdf(
    input_path: str, 
    output_folder: str, 
    mode: str = "single",  # "single" (一頁一檔案), "range" (多頁合成一檔)
    ranges: str = "",      
    custom_name: str = "",
    callback: Callable[[float], None] | None = None
) -> tuple[bool, str]:
    """
    拆分 PDF 檔案。
    """
    try:
        reader = pypdf.PdfReader(input_path)
        total_pages = len(reader.pages)
        base_name = custom_name if custom_name else os.path.splitext(os.path.basename(input_path))[0]

        if not os.path.exists(output_folder):
            os.makedirs(output_folder, exist_ok=True)

        target_indices = parse_range_string(ranges, total_pages)
        if not target_indices:
            return False, "未偵測到有效的頁碼範圍。"

        if mode == "single":
            # 一頁一檔案
            for i, p_idx in enumerate(target_indices):
                writer = pypdf.PdfWriter()
                writer.add_page(reader.pages[p_idx])
                output_filename = f"{base_name}_page_{p_idx+1:03d}.pdf"
                with open(os.path.join(output_folder, output_filename), "wb") as f:
                    writer.write(f)
                
                if callback:
                    callback((i + 1) / len(target_indices))
            return True, f"成功！已將選取的 {len(target_indices)} 頁分別儲存為單一檔案。"

        elif mode == "range":
            # 多頁合成一檔
            writer = pypdf.PdfWriter()
            for i, p_idx in enumerate(target_indices):
                writer.add_page(reader.pages[p_idx])
                if callback:
                    callback((i + 1) / len(target_indices))
            
            output_filename = f"{base_name}_extracted.pdf"
            with open(os.path.join(output_folder, output_filename), "wb") as f:
                writer.write(f)
            
            return True, f"成功！已將選取的 {len(target_indices)} 頁合併提取至 {output_filename}"

    except Exception as e:
        return False, f"拆分失敗：\n{str(e)}"

def pdf_to_jpg(
    input_path: str, 
    output_folder: str, 
    dpi: int = 200, 
    quality: int = 85, 
    ranges: str = "", 
    custom_name: str = "",
    callback: Callable[[float], None] | None = None
) -> tuple[bool, str]:
    """
    將 PDF 轉換為一系列 JPG 圖片 (使用 PyMuPDF，無需外部 Poppler 依賴)。
    """
    try:
        doc = fitz.open(input_path)
        total_pages = len(doc)
        
        if total_pages == 0:
            doc.close()
            return False, "該 PDF 檔案沒有任何頁面。"

        if not os.path.exists(output_folder):
            os.makedirs(output_folder, exist_ok=True)
            
        base_name = custom_name if custom_name else os.path.splitext(os.path.basename(input_path))[0]
        
        target_indices = parse_range_string(ranges, total_pages)
        if not target_indices:
            doc.close()
            return False, "未偵測到有效的頁碼範圍。"

        # 計算 DPI 縮放矩陣 (PDF 預設基準為 72 DPI)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)

        for i, p_idx in enumerate(target_indices):
            page_num = p_idx + 1
            page = doc.load_page(p_idx)
            # 渲染為 RGB 影像
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            output_filename = f"{base_name}_page_{page_num:03d}.jpg"
            save_path = os.path.join(output_folder, output_filename)
            
            img_bytes = pix.tobytes("jpeg", jpg_quality=quality)
            with open(save_path, "wb") as f:
                f.write(img_bytes)
            
            if callback:
                callback((i + 1) / len(target_indices))
                
        doc.close()
        return True, f"成功！已將選取的 {len(target_indices)} 頁轉換為圖片並儲存至 {output_folder}"

    except Exception as e:
        return False, f"轉換過程中發生錯誤：\n{str(e)}"

def compress_pdf(
    input_path: str,
    output_folder: str,
    quality: str = "medium", # "low", "medium", "high"
    custom_name: str = "",
    callback: Callable[[float], None] | None = None
) -> tuple[bool, str]:
    """
    壓縮 PDF 檔案大小。
    """
    try:
        reader = pypdf.PdfReader(input_path)
        writer = pypdf.PdfWriter()
        
        base_name = custom_name if custom_name else os.path.splitext(os.path.basename(input_path))[0]
        output_filename = f"{base_name}_compressed.pdf"
        output_path = os.path.join(output_folder, output_filename)

        if not os.path.exists(output_folder):
            os.makedirs(output_folder, exist_ok=True)

        # 根據等級決定圖片品質與縮放限制
        img_quality = 65 if quality == "low" else (40 if quality == "medium" else 20)
        max_dim = 1800 if quality == "low" else (1200 if quality == "medium" else 800)

        processed_ids = set()

        total_pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            new_page = writer.add_page(page)
            new_page.compress_content_streams() # 壓縮文字與路徑
            
            try:
                for img_proxy in new_page.images:
                    try:
                        img_id = img_proxy.indirect_reference.idnum
                    except:
                        continue 

                    if img_id not in processed_ids:
                        # 取得原始資料大小
                        original_data_size = len(img_proxy.data)
                        pil_img = img_proxy.image
                        
                        # 轉為 RGB 以確保 JPEG 壓縮效率 (移除透明通道)
                        if pil_img.mode in ("RGBA", "P"):
                            pil_img = pil_img.convert("RGB")
                        
                        w, h = pil_img.size
                        if max(w, h) > max_dim:
                            ratio = max_dim / max(w, h)
                            pil_img = pil_img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
                        
                        # 在記憶體中模擬壓縮後的結果
                        img_byte_arr = io.BytesIO()
                        pil_img.save(img_byte_arr, format="JPEG", quality=img_quality)
                        compressed_data_size = img_byte_arr.tell()

                        # 終極保換：只有當壓縮後真的變小，才進行替換
                        if compressed_data_size < original_data_size:
                            img_proxy.replace(pil_img, quality=img_quality)
                        
                        processed_ids.add(img_id)
            except Exception:
                pass
            
            if callback:
                callback((i + 1) / total_pages)

        # 寫入檔案
        with open(output_path, "wb") as f:
            writer.write(f)

        return True, output_path
    except Exception as e:
        return False, f"壓縮失敗: {str(e)}"

def get_installed_ocr_languages() -> list[str]:
    """
    獲取系統已安裝的 Windows OCR 語言標籤。
    """
    import subprocess
    import os
    try:
        ps_cmd = (
            "[void][System.Reflection.Assembly]::LoadWithPartialName('System.Runtime.WindowsRuntime'); "
            "[void][Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType = WindowsRuntime]; "
            "[Windows.Media.Ocr.OcrEngine]::AvailableRecognizerLanguages | Select-Object -ExpandProperty LanguageTag"
        )
        cmd = ["powershell", "-ExecutionPolicy", "Bypass", "-Command", ps_cmd]
        
        startupinfo = None
        if os.name == 'nt':
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            startupinfo.wShowWindow = subprocess.SW_HIDE
            
        res = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', errors='ignore', startupinfo=startupinfo)
        tags = [t.strip() for t in res.stdout.split('\n') if t.strip()]
        return tags
    except:
        return []

def convert_pdf_to_word(
    input_path: str,
    output_folder: str,
    mode: str = "digital",  # "digital" or "ocr"
    custom_name: str = "",
    callback: Callable[[float, str], None] | None = None
) -> tuple[bool, str]:
    """
    將 PDF 轉換為 Word (.docx) 檔案。
    """
    try:
        import os
        base_name = custom_name if custom_name else os.path.splitext(os.path.basename(input_path))[0]
        output_filename = f"{base_name}.docx"
        output_path = os.path.join(output_folder, output_filename)

        if not os.path.exists(output_folder):
            os.makedirs(output_folder, exist_ok=True)

        if mode == "digital":
            from pdf2docx import Converter
            cv = Converter(input_path)
            total_pages = len(cv.fitz_doc)
            
            if total_pages == 0:
                cv.close()
                return False, "該 PDF 檔案沒有任何頁面。"

            # 1. 載入所有頁面
            cv.load_pages(0, total_pages, None)
            
            # 2. 解析文件層級結構
            settings = cv.default_settings
            cv.parse_document(**settings)
            
            # 3. 逐頁進行佈局分析以更新進度
            for i, page in enumerate(cv.pages):
                if page.skip_parsing:
                    continue
                if callback:
                    callback(i / total_pages, f"正在分析第 {i+1} 頁佈局 (共 {total_pages} 頁)...")
                page.parse(**settings)
            
            if callback:
                callback(0.95, "正在產生 Word 檔案...")
                
            # 4. 生成並儲存 docx (只存檔一次，保留所有頁面與表格)
            cv.make_docx(output_path, **settings)
            cv.close()
            return True, f"轉換成功！Word 檔案已儲存至：\n{output_path}"

        elif mode == "ocr":
            import docx
            import subprocess
            import tempfile
            import time

            # 偵測是否安裝中文 OCR
            langs = get_installed_ocr_languages()
            has_chinese = any(lang.lower().startswith("zh") for lang in langs)
            if not has_chinese:
                if callback:
                    callback(0.0, "⚠️ 系統未安裝中文 OCR 語言包，辨識結果可能為英文/亂碼...")
                time.sleep(2.5)

            # 1. 取得 PDF 資訊
            fitz_doc = fitz.open(input_path)
            total_pages = len(fitz_doc)
            if total_pages == 0:
                fitz_doc.close()
                return False, "該 PDF 檔案沒有任何頁面。"

            doc = docx.Document()
            zoom = 150 / 72.0
            mat = fitz.Matrix(zoom, zoom)
            
            # 在臨時目錄中建立 OCR 腳本與影像
            with tempfile.TemporaryDirectory() as temp_dir:
                # 寫入 PowerShell OCR 腳本
                ps_script_path = os.path.join(temp_dir, "ocr.ps1")
                ps_script_content = """param (
    [string]$ImagePath
)
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
try {
    Add-Type -AssemblyName System.Runtime.WindowsRuntime
    [void][Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType = WindowsRuntime]
    [void][Windows.Graphics.Imaging.SoftwareBitmap, Windows.Foundation, ContentType = WindowsRuntime]
    [void][Windows.Storage.StorageFile, Windows.Foundation, ContentType = WindowsRuntime]
    [void][Windows.Storage.Streams.IRandomAccessStream, Windows.Foundation, ContentType = WindowsRuntime]

    $asTaskGeneric = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object { 
        $_.Name -eq 'AsTask' -and 
        $_.GetParameters().Count -eq 1 -and 
        $_.GetParameters()[0].ParameterType.Name -eq 'IAsyncOperation`1' 
    })[0]

    function Await-WinRT($asyncOperation, $resultType) {
        $asTask = $asTaskGeneric.MakeGenericMethod($resultType)
        $netTask = $asTask.Invoke($null, @($asyncOperation))
        $netTask.Wait(-1) | Out-Null
        return $netTask.Result
    }

    $absPath = [System.IO.Path]::GetFullPath($ImagePath)
    $fileOperation = [Windows.Storage.StorageFile]::GetFileFromPathAsync($absPath)
    $file = Await-WinRT $fileOperation ([Windows.Storage.StorageFile])

    $streamOperation = $file.OpenAsync([Windows.Storage.FileAccessMode]::Read)
    $stream = Await-WinRT $streamOperation ([Windows.Storage.Streams.IRandomAccessStream])

    $decoderOperation = [Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)
    $decoder = Await-WinRT $decoderOperation ([Windows.Graphics.Imaging.BitmapDecoder])

    $bitmapOperation = $decoder.GetSoftwareBitmapAsync()
    $softwareBitmap = Await-WinRT $bitmapOperation ([Windows.Graphics.Imaging.SoftwareBitmap])

    [void][Windows.Globalization.Language, Windows.Foundation, ContentType = WindowsRuntime]
    $lang = New-Object Windows.Globalization.Language("zh-Hant-TW")
    $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage($lang)
    if ($null -eq $engine) {
        $lang = New-Object Windows.Globalization.Language("zh-Hans-CN")
        $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage($lang)
    }
    if ($null -eq $engine) {
        $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
    }
    if ($null -eq $engine) {
        exit 1
    }

    $ocrOperation = $engine.RecognizeAsync($softwareBitmap)
    $result = Await-WinRT $ocrOperation ([Windows.Media.Ocr.OcrResult])

    Write-Output $result.Text
}
catch {
    exit 1
}
"""
                with open(ps_script_path, "w", encoding="utf-8") as f:
                    f.write(ps_script_content)

                # 逐頁轉圖並執行 OCR
                for i in range(total_pages):
                    page_num = i + 1
                    
                    if callback:
                        callback(i / total_pages, f"正在將第 {page_num} 頁轉換為圖片...")
                    
                    page = fitz_doc.load_page(i)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    temp_png = os.path.join(temp_dir, f"page_{page_num}.png")
                    pix.save(temp_png)
                    
                    if callback:
                        callback((i + 0.5) / total_pages, f"正在辨識第 {page_num} 頁文字 (OCR)...")
                    
                    # 執行 PowerShell OCR
                    cmd = [
                        "powershell",
                        "-ExecutionPolicy", "Bypass",
                        "-File", ps_script_path,
                        "-ImagePath", temp_png
                    ]
                    
                    # Windows 底下隱藏視窗
                    startupinfo = None
                    if os.name == 'nt':
                        startupinfo = subprocess.STARTUPINFO()
                        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                        startupinfo.wShowWindow = subprocess.SW_HIDE
                    
                    res = subprocess.run(
                        cmd, 
                        capture_output=True, 
                        text=True, 
                        encoding='utf-8', 
                        errors='ignore',
                        startupinfo=startupinfo
                    )
                    
                    ocr_text = res.stdout.strip()
                    
                    # 寫入 Word
                    if i > 0:
                        doc.add_page_break()
                    
                    doc.add_heading(f"Page {page_num}", level=2)
                    if ocr_text:
                        # 依換行分割段落寫入
                        for line in ocr_text.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("[此頁未偵測到文字]")
                    
                    if callback:
                        callback((i + 1) / total_pages, f"第 {page_num} 頁處理完成")
                
            fitz_doc.close()
            if callback:
                callback(1.0, "儲存 Word 檔案中...")
            doc.save(output_path)
            return True, f"OCR 辨識轉換成功！Word 檔案已儲存至：\n{output_path}"

    except Exception as e:
        return False, f"轉換失敗: {str(e)}"

# ── 頁面編輯類功能 ─────────────────────────────────────────────

def generate_page_thumbnail(path: str, page_idx: int, rotation: int = 0):
    """
    生成指定頁面的縮圖。
    如果是 PDF 且 idx >= 0，則提取頁面；如果是圖片，則直接讀取。
    """
    from PIL import Image
    import io
    
    try:
        if path.lower().endswith((".jpg", ".png", ".jpeg")):
            img = Image.open(path)
        else:
            doc = fitz.open(path)
            page = doc.load_page(page_idx)
            pix = page.get_pixmap(alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            doc.close()
        
        # 套用旋轉
        if rotation != 0:
            img = img.rotate(-rotation, expand=True) # PIL 旋轉是逆時針，所以加負號
            
        return img
    except:
        return None

def save_manipulated_pdf(pages_data: list, output_path: str):
    """
    根據編輯數據儲存 PDF。
    data: list of {"path": str, "idx": int, "rotation": int}
    """
    import pypdf
    from PIL import Image
    import io
    import os

    writer = pypdf.PdfWriter()
    readers = {} # 快取 Reader

    try:
        for item in pages_data:
            path = item["path"]
            idx = item["idx"]
            rot = item["rotation"]
            
            if path.lower().endswith(".pdf") and idx >= 0:
                if path not in readers:
                    readers[path] = pypdf.PdfReader(path)
                
                page = readers[path].pages[idx]
                if rot != 0:
                    page.add_transformation(pypdf.Transformation().rotate(rot))
                writer.add_page(page)
            else:
                # 處理圖片插入
                img = Image.open(path).convert("RGB")
                pdf_bytes = io.BytesIO()
                img.save(pdf_bytes, format="PDF")
                img_reader = pypdf.PdfReader(pdf_bytes)
                page = img_reader.pages[0]
                if rot != 0:
                    page.add_transformation(pypdf.Transformation().rotate(rot))
                writer.add_page(page)
        
        with open(output_path, "wb") as f:
            writer.write(f)
        return output_path
    except Exception as e:
        import traceback
        traceback.print_exc()
        return str(e)
